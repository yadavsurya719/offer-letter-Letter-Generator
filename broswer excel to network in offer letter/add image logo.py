import tkinter as tk
from tkinter import filedialog
from docx import Document
from datetime import datetime
from docx2pdf import convert
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PIL import Image


def process_excel(file_path, text_widget):
    try:
        df = pd.read_excel(file_path)

        # Clear previous content in the Text widget
        text_widget.delete(1.0, tk.END)

        # Create leave letters for all employees
        for index, row in df.iterrows():
            leave_letter = offer_letter(row['employee_name'], str(row['company_name']), str(row['job_name']),str(row['start_date']),str(row['salary']), row['hr_name'])
            text_widget.insert(tk.END, leave_letter + '\n')

    except Exception as e:
        text_widget.insert(tk.END, f"Error processing Excel file: {e}\n")

def browse_excel_file(entry_path):
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    entry_path.delete(0, tk.END)
    entry_path.insert(0, file_path)

def process_button_click(entry_path, text_widget):
    file_path = entry_path.get()
    if file_path:
        process_excel(file_path, text_widget)
    else:
        text_widget.insert(tk.END, "Please select an Excel file.\n")


        
def offer_letter(employee_name,company_name,job_name ,start_date,salary,hr_name):

    pdf_filename = f'offer_Letter_{employee_name.replace(" ","_")}.pdf'
    c = canvas.Canvas(pdf_filename, pagesize=letter)

    # Add image as background
    image_path = r'C:\surya\vs code\broswer excel to network in offer letter\1588848.jpg'  # Replace with the path to your image file
    c.drawImage(image_path, 0, 0, width=letter[0], height=letter[1])

    
    # Your existing text content
    text_content = offer_letter(employee_name, company_name, job_name, start_date, salary, hr_name)

    # Add text on top of the image
    c.setFont("Helvetica", 12)
    c.drawString(100, 500, text_content)

    # Save the PDF
    c.save()

    return pdf_filename
    # Create a new Document
    doc = Document()

    # Add a heading
    doc.add_heading('Offer letter', level=1)


    # Add the current date
    current_date = datetime.now().strftime('%Y-%m-%d')
    doc.add_paragraph(f'Date: {current_date}')

    # Add employee details
    doc.add_paragraph(f'Dear {employee_name},')
    #doc.add_paragraph(f'Job Role : {job_name}')
    #doc.add_paragraph(f'Salary: {salary}')
    #doc.add_paragraph(f'HR : {hr_name}')
    doc.add_paragraph(f'We at {company_name} are delighted to inform you that you have been selected to our company as a {job_name}. Your joining date is {start_date} and you have to report to the office at 9:AM.Your salary willbe {salary} per month.')
    #doc.add_paragraph(f' {'name'}')
    #doc.add_paragraph(f'Employee Name: {employee_name}')
    #doc.add_paragraph(f'Company: {company_name}')
    #doc.add_paragraph(f'Job Role : {job_name}')
    #doc.add_paragraph(f'Joining date: {start_date}')
    #doc.add_paragraph(f'Salary: {salary}')
    

    # Extract the date part if time information is present
    #start_date = str(start_date).split()[0]
    #end_date = str(end_date).split()[0]


 # Calculate the number of days
    #start_date_dt = datetime.strptime(start_date, '%Y-%m-%d')

    #end_date_dt = datetime.strptime(end_date, '%Y-%m-%d')
    #num_days = (end_date_dt - start_date_dt).days + 1

    #doc.add_paragraph(f'Total Days: {num_days}')

    # Add a closing statement
    doc.add_paragraph(f'Attached here to is your offer letter. If you accept our offer,please make yoyrself available on the date mentioned in the document for your official first date of reporting. You are required to bring the signed copy of the offer letter as well as the necessary documents whose list is also attached below . \n\n Welcome to {company_name}....We are looking farward to working with you.')
    #doc.add_paragraph(f'Company : {company_name}')
    
    doc.add_paragraph('Regards,')
    # Add a closing salutation
    #doc.add_paragraph(f'HR : {hr_name}')
    doc.add_paragraph(f'{hr_name}[HR], {company_name} . ')

    # Save the document
    doc_filename = f'offer_Letter_{employee_name.replace(" ","_")}.docx'
    doc.save(doc_filename)

    # Convert the document to PDF
    convert(doc_filename)

    
    return f'Leave letter created for {employee_name}. Saved as {doc_filename} and converted to PDF.'

# Create the main application window
app = tk.Tk()
app.title("offer Letter Generator")

# Create widgets
label_path = tk.Label(app, text="Enter Excel File Path:")
entry_path = tk.Entry(app, width=40)
button_browse = tk.Button(app, text="Browse", command=lambda: browse_excel_file(entry_path))
button_process = tk.Button(app, text="Generate offer Letters", command=lambda: process_button_click(entry_path, text_widget))

# Text widget to display the generated leave letters
text_widget = tk.Text(app, height=10, width=80)

# Place widgets in the window
label_path.grid(row=0, column=0, padx=10, pady=10)
entry_path.grid(row=0, column=1, padx=10, pady=10)
button_browse.grid(row=0, column=2, padx=10, pady=10)
button_process.grid(row=1, column=0, columnspan=3, pady=10)
text_widget.grid(row=2, column=0, columnspan=3, pady=10)

# Run the application
app.mainloop()
